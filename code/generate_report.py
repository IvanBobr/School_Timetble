from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_TITLE = (
    "Создание автоматической системы визуального отображения информации о расписании уроков и его изменениях"
)

# Данные, которые пользователь уже предоставил
STUDENT_FIO = "Бобринев Иван Юрьевич"
STUDENT_GROUP = "М8О-115БВ-25"
ORGANIZATION_FULL = "ФГАОУ ВО «Московский авиационный институт (национальный исследовательский университет)» (МАИ)"
SUPERVISOR_ORDER = "доцент каф. ОЦ 8 Пановский Валентин Николаевич"


@dataclass(frozen=True)
class Paths:
    root: Path
    readme: Path
    requirements: Path
    out_docx: Path


def red_placeholder(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    # Ensure East Asia font too (Word sometimes ignores font.name)
    rpr = style.element.rPr
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rfonts)


def set_page_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n").bold = True
    p.add_run(ORGANIZATION_FULL).bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА\n").bold = True
    p2.add_run("на тему:\n")
    p2.add_run(PROJECT_TITLE).bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.add_run("Выполнил(а): ").bold = True
    p3.add_run(f"{STUDENT_FIO}, {STUDENT_GROUP}\n")
    p3.add_run("Руководитель: ").bold = True
    p3.add_run(f"{SUPERVISOR_ORDER}\n")

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    red_placeholder(p4, "[ЗАПОЛНИТЕ: город] ")
    p4.add_run(str(datetime.now().year))

    doc.add_page_break()


def add_executors_page(doc: Document) -> None:
    p = doc.add_paragraph("СПИСОК ИСПОЛНИТЕЛЕЙ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.add_run(f"{STUDENT_FIO} – выполнены все разделы выпускной квалификационной работы.")
    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    p = doc.add_paragraph("РЕФЕРАТ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.add_run("Выпускная квалификационная работа состоит из ")
    red_placeholder(p2, "[ЗАПОЛНИТЕ: __] ")
    p2.add_run("страниц, ")
    red_placeholder(p2, "[ЗАПОЛНИТЕ: __] ")
    p2.add_run("рисунков, ")
    red_placeholder(p2, "[ЗАПОЛНИТЕ: __] ")
    p2.add_run("таблиц, ")
    red_placeholder(p2, "[ЗАПОЛНИТЕ: __] ")
    p2.add_run("использованных источников, ")
    red_placeholder(p2, "[ЗАПОЛНИТЕ: __] ")
    p2.add_run("приложений.")

    p3 = doc.add_paragraph()
    p3.paragraph_format.first_line_indent = Cm(1.25)
    p3.add_run("Ключевые слова: ")
    p3.add_run(
        "РАСПИСАНИЕ, ШКОЛА, ИЗМЕНЕНИЯ РАСПИСАНИЯ, ЛОКАЛЬНАЯ СЕТЬ, КЛИЕНТ-СЕРВЕР, "
        "PYTHON, FLASK, TKINTER, JSON, КЭШИРОВАНИЕ, АВТОМАТИЗАЦИЯ"
    )

    p4 = doc.add_paragraph()
    p4.paragraph_format.first_line_indent = Cm(1.25)
    p4.add_run(
        "Итоговая аттестационная работа выполнена в формате IT-проекта и посвящена разработке "
        "автоматизированной системы визуального отображения школьного расписания и его изменений в "
        "условиях работы в локальной сети. Система включает серверный компонент, формирующий и "
        "раздающий актуальные данные, и клиентское приложение, отображающее расписание и применяющее "
        "изменения на основе предоставленных администрацией документов."
    )

    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    p = doc.add_paragraph("СОДЕРЖАНИЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p2 = doc.add_paragraph()
    red_placeholder(
        p2,
        "[ЗАПОЛНИТЕ: содержание (оглавление) должно быть оформлено с точками-лидерами и номерами страниц. "
        "Можно сформировать автоматически в Word через «Ссылки → Оглавление» после применения стилей заголовков.]\n",
    )
    doc.add_page_break()


def add_terms_and_definitions(doc: Document) -> None:
    p = doc.add_paragraph("ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.add_run(
        "В настоящей итоговой аттестационной работе применяют следующие термины с соответствующими определениями:"
    )

    terms = [
        ("Клиент-серверная архитектура", "модель взаимодействия, в которой клиент запрашивает данные у сервера по сети."),
        ("Кэш", "локально сохраненные данные, используемые при временной недоступности сервера."),
        ("Изменения расписания", "корректировки базового расписания (замены и отмены уроков) на определенную дату."),
        ("JSON", "текстовый формат обмена данными, используемый для передачи расписания от сервера к клиенту."),
    ]
    for term, definition in terms:
        p2 = doc.add_paragraph()
        p2.add_run(f"{term} – {definition}")

    doc.add_page_break()


def add_abbreviations(doc: Document) -> None:
    p = doc.add_paragraph("ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.add_run(
        "В настоящей итоговой аттестационной работе применяют следующие сокращения и обозначения:"
    )

    items = [
        ("ЛВС", "локальная вычислительная сеть"),
        ("ПО", "программное обеспечение"),
        ("API", "программный интерфейс приложения"),
        ("GUI", "графический интерфейс пользователя"),
    ]
    for abbr, meaning in items:
        p2 = doc.add_paragraph()
        p2.add_run(f"{abbr} – {meaning}")

    doc.add_page_break()


def add_intro(doc: Document) -> None:
    p = doc.add_paragraph("ВВЕДЕНИЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(
        "Актуальность темы обусловлена необходимостью оперативного доведения до учащихся и педагогов "
        "информации о расписании и его изменениях. При ручном обновлении расписания увеличивается риск "
        "ошибок и задержек, особенно при частых заменах и отменах уроков. Автоматизация процесса позволяет "
        "сократить трудозатраты и повысить достоверность отображаемой информации."
    )

    par2 = doc.add_paragraph()
    par2.paragraph_format.first_line_indent = Cm(1.25)
    par2.add_run("Цель работы – разработать автоматизированную систему визуального отображения школьного расписания и его изменений в локальной сети.")

    par3 = doc.add_paragraph()
    par3.paragraph_format.first_line_indent = Cm(1.25)
    par3.add_run("Для достижения цели решены следующие задачи:")
    ul = [
        "проанализировать предметную область и существующие подходы к хранению/распространению расписаний;",
        "разработать клиент-серверную архитектуру обмена данными в ЛВС;",
        "реализовать парсинг базового расписания из файла Excel и изменений из файла Word;",
        "реализовать серверный API для передачи актуального расписания клиентам;",
        "реализовать клиентское приложение с визуальным отображением расписания и применением изменений;",
        "обеспечить работу при недоступности сервера за счет локального кэша.",
    ]
    for i, item in enumerate(ul):
        pp = doc.add_paragraph(style="List Bullet")
        pp.add_run(item)

    par4 = doc.add_paragraph()
    par4.paragraph_format.first_line_indent = Cm(1.25)
    par4.add_run("Объект разработки – информационная система отображения расписания уроков в школе.")

    par5 = doc.add_paragraph()
    par5.paragraph_format.first_line_indent = Cm(1.25)
    par5.add_run("Предмет разработки – программные методы получения, обработки и визуализации расписания и его изменений, а также сетевое распространение данных.")


def add_chapter_1_theory(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("1 ТЕОРЕТИЧЕСКАЯ ЧАСТЬ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    s1 = doc.add_paragraph("1.1 Автоматизация информирования о расписании")
    s1.runs[0].bold = True
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(
        "В школьной среде расписание подвержено изменениям (замены, переносы, отмены). "
        "Эффективная система должна обеспечивать единый источник актуальных данных и быстрое распространение "
        "информации на устройства отображения (мониторы, ПК в кабинетах, информационные панели)."
    )

    s2 = doc.add_paragraph("1.2 Клиент-серверный подход в локальной сети")
    s2.runs[0].bold = True
    par2 = doc.add_paragraph()
    par2.paragraph_format.first_line_indent = Cm(1.25)
    par2.add_run(
        "Клиент-серверная архитектура позволяет централизованно формировать данные на сервере и "
        "предоставлять их клиентам через HTTP API. Это упрощает обновление, обеспечивает согласованность "
        "данных и дает возможность масштабировать число клиентов без усложнения логики распространения."
    )

    s3 = doc.add_paragraph("1.3 Форматы хранения и обмена данными")
    s3.runs[0].bold = True
    par3 = doc.add_paragraph()
    par3.paragraph_format.first_line_indent = Cm(1.25)
    par3.add_run(
        "В качестве исходных форматов используются документы, привычные для администрации: Excel "
        "для базового расписания и Word для перечня изменений. Для сетевого обмена используется JSON, "
        "который удобен для сериализации структурированных данных и поддерживается большинством языков программирования."
    )


def add_chapter_2_practice(doc: Document, readme_text: str) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("2 ПРАКТИЧЕСКАЯ ЧАСТЬ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    s1 = doc.add_paragraph("2.1 Требования и исходные данные")
    s1.runs[0].bold = True
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(
        "Система предназначена для работы в общей школьной сети и обеспечивает раздачу актуальных данных "
        "с одного хоста (сервер) на остальные устройства (клиенты). Входными данными являются файлы "
        "`rasp.xls` (базовое расписание) и `changes.docx` (изменения на дату)."
    )

    par2 = doc.add_paragraph()
    par2.paragraph_format.first_line_indent = Cm(1.25)
    par2.add_run("Краткое описание проекта (из README):")

    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Cm(1.25)
    q.paragraph_format.right_indent = Cm(1.25)
    q.add_run(readme_text.strip()[:900] + ("…" if len(readme_text.strip()) > 900 else ""))

    s2 = doc.add_paragraph("2.2 Архитектура программного решения")
    s2.runs[0].bold = True
    par3 = doc.add_paragraph()
    par3.paragraph_format.first_line_indent = Cm(1.25)
    par3.add_run(
        "Решение состоит из серверного приложения (Flask) и клиентского приложения (Tkinter). "
        "Сервер формирует структуру данных расписания, объединяя базовое расписание и список изменений, "
        "и предоставляет ее по HTTP. Клиент запрашивает данные у сервера, сохраняет кэш на диск и отображает "
        "расписание в нескольких режимах (все классы/один класс/все дни)."
    )

    s3 = doc.add_paragraph("2.3 Серверный модуль")
    s3.runs[0].bold = True
    par4 = doc.add_paragraph()
    par4.paragraph_format.first_line_indent = Cm(1.25)
    par4.add_run(
        "Сервер реализован в модуле `code/server.py` и запускается на адресе 0.0.0.0:5000. "
        "Основная точка доступа `GET /api/schedule` возвращает JSON с расписанием. "
        "Формирование JSON выполняется функцией `make_jsonFile()` из `code/data_to_jsonFile.py`."
    )

    s4 = doc.add_paragraph("2.4 Формирование данных (Excel + Word)")
    s4.runs[0].bold = True
    par5 = doc.add_paragraph()
    par5.paragraph_format.first_line_indent = Cm(1.25)
    par5.add_run(
        "Модуль `code/parsing_excel.py` извлекает список классов, кабинетов и предметов из Excel, а также "
        "строит расписание по дням недели. Модуль `code/parsing_word.py` извлекает изменения из Word-документа: "
        "замены уроков и уроки, которые нужно пропустить. Эти данные объединяются в единую структуру JSON."
    )

    s5 = doc.add_paragraph("2.5 Клиентский модуль и визуализация")
    s5.runs[0].bold = True
    par6 = doc.add_paragraph()
    par6.paragraph_format.first_line_indent = Cm(1.25)
    par6.add_run(
        "Клиент реализован в `code/timetable.py` (GUI на Tkinter). При запуске клиент запрашивает JSON с сервера. "
        "При успешном ответе данные сохраняются в `cached_schedule.json`. Если сервер недоступен, данные загружаются "
        "из кэша, что обеспечивает устойчивую работу в сети. Клиент применяет изменения к расписанию на текущую дату "
        "и подсвечивает статусы (по расписанию/изменено/отменено)."
    )

    s6 = doc.add_paragraph("2.6 Алгоритм применения изменений")
    s6.runs[0].bold = True
    par7 = doc.add_paragraph()
    par7.paragraph_format.first_line_indent = Cm(1.25)
    par7.add_run(
        "Изменения применяются только если дата в `changes.docx` совпадает с текущей датой. Для замен "
        "соответствующий урок в расписании класса заменяется на запись с пометкой «ИЗМЕНЕНО». "
        "Для пропусков урок помечается как «ОТМЕНЕНО» и очищаются поля предмета/кабинета/учителя."
    )

    s7 = doc.add_paragraph("2.7 Тестирование и результаты")
    s7.runs[0].bold = True
    par8 = doc.add_paragraph()
    par8.paragraph_format.first_line_indent = Cm(1.25)
    red_placeholder(
        par8,
        "[ЗАПОЛНИТЕ: опишите, как вы тестировали систему (сценарии: сервер доступен, сервер недоступен, "
        "изменения на сегодня/на другую дату), и какие результаты получили. Добавьте скриншоты в приложение, "
        "а здесь сделайте ссылки на рисунки.]\n",
    )


def add_user_guide(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("3 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    s1 = doc.add_paragraph("3.1 Подготовка файлов расписания")
    s1.runs[0].bold = True
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(
        "На компьютере-хосте (сервер) в папку `downloads/` рядом с серверным приложением помещаются файлы "
        "`rasp.xls` и `changes.docx`. Файл `rasp.xls` содержит базовое расписание, а `changes.docx` — перечень изменений."
    )

    s2 = doc.add_paragraph("3.2 Запуск сервера")
    s2.runs[0].bold = True
    par2 = doc.add_paragraph()
    par2.paragraph_format.first_line_indent = Cm(1.25)
    par2.add_run("Сервер запускается командой:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(1.25)
    code.add_run("python code/server.py")

    s3 = doc.add_paragraph("3.3 Запуск клиента")
    s3.runs[0].bold = True
    par3 = doc.add_paragraph()
    par3.paragraph_format.first_line_indent = Cm(1.25)
    par3.add_run(
        "Клиент запускается на устройствах отображения. По умолчанию клиент запрашивает сервер по адресу "
        "127.0.0.1 (если клиент и сервер на одном ПК). Для работы по сети необходимо указать IP-адрес хоста."
    )
    par4 = doc.add_paragraph()
    par4.paragraph_format.first_line_indent = Cm(1.25)
    red_placeholder(
        par4,
        "[ЗАПОЛНИТЕ: укажите реальный IP сервера в вашей сети и как вы его передаете клиенту "
        "(например, правкой `server_ip` в коде или отдельной настройкой).]\n",
    )

    s4 = doc.add_paragraph("3.4 Основные режимы работы")
    s4.runs[0].bold = True
    items = [
        "просмотр текущего или следующего урока для всех классов;",
        "просмотр расписания конкретного класса на текущий день (все уроки или только будущие);",
        "просмотр полного расписания по дням недели и по группам (параллелям).",
    ]
    for it in items:
        pp = doc.add_paragraph(style="List Bullet")
        pp.add_run(it)


def add_conclusion(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(
        "В ходе выполнения выпускной квалификационной работы разработана автоматизированная система "
        "визуального отображения школьного расписания и его изменений для работы в локальной сети. "
        "Реализованы серверный компонент для формирования и передачи актуального расписания, клиентское приложение "
        "для отображения данных и механизм кэширования на случай недоступности сервера."
    )

    par2 = doc.add_paragraph()
    par2.paragraph_format.first_line_indent = Cm(1.25)
    red_placeholder(
        par2,
        "[ЗАПОЛНИТЕ: перечислите конкретные результаты (что именно работает), ограничения (например, формат входных файлов), "
        "и направления дальнейшего развития (например, конфигурация IP без правки кода, авторизация, веб-интерфейс, автопоиск сервера).]\n",
    )


def add_sources(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    access = date.today().strftime("%d.%m.%Y")
    sources = [
        # Нормативка / стандарты
        (
            "ГОСТ 7.32–2017. Система стандартов по информации, библиотечному и издательскому делу. "
            "Отчет о научно-исследовательской работе. Структура и правила оформления. "
            "Электронный ресурс. – Режим доступа: https://base.garant.ru/71884728/ (дата обращения: "
            f"{access})."
        ),
        (
            "ГОСТ 7.32–2017 (PDF). Отчет о научно-исследовательской работе. Структура и правила оформления. "
            "Электронный ресурс. – Режим доступа: https://science.itmo.ru/wp-content/uploads/2022/08/gost_7.32_2017.pdf "
            f"(дата обращения: {access})."
        ),
        # Язык/GUI
        (
            "Python Software Foundation. The Python Standard Library (Python 3.14). "
            "Электронный ресурс. – Режим доступа: https://docs.python.org/3/library/index.html "
            f"(дата обращения: {access})."
        ),
        (
            "Python Software Foundation. tkinter — Python interface to Tcl/Tk (Python 3.14). "
            "Электронный ресурс. – Режим доступа: https://docs.python.org/3/library/tkinter.html "
            f"(дата обращения: {access})."
        ),
        # Сервер/HTTP
        (
            "Pallets Projects. Flask Documentation (Flask 3.1.x). Электронный ресурс. – Режим доступа: "
            "https://flask.palletsprojects.com/en/stable/ (дата обращения: "
            f"{access})."
        ),
        (
            "Requests. Requests: HTTP for Humans (v2.34.x). Электронный ресурс. – Режим доступа: "
            "https://docs.python-requests.org/en/latest/ (дата обращения: "
            f"{access})."
        ),
        # Обработка данных/файлов
        (
            "pandas development team. pandas documentation (Version 3.0.x). Электронный ресурс. – Режим доступа: "
            "https://pandas.pydata.org/docs/index.html (дата обращения: "
            f"{access})."
        ),
        (
            "openpyxl documentation. openpyxl — A Python library to read/write Excel 2010 xlsx/xlsm files. "
            "Электронный ресурс. – Режим доступа: https://openpyxl.readthedocs.io/en/stable/ "
            f"(дата обращения: {access})."
        ),
        (
            "python-docx documentation. python-docx — Python library for creating and updating Microsoft Word (.docx) files. "
            "Электронный ресурс. – Режим доступа: https://python-docx.readthedocs.io/en/latest/index.html "
            f"(дата обращения: {access})."
        ),
        # Формат обмена
        (
            "RFC 8259. The JavaScript Object Notation (JSON) Data Interchange Format (STD 90). "
            "Электронный ресурс. – Режим доступа: https://www.rfc-editor.org/info/rfc8259/ "
            f"(дата обращения: {access})."
        ),
    ]

    for i, s in enumerate(sources, start=1):
        par = doc.add_paragraph()
        par.paragraph_format.first_line_indent = Cm(1.25)
        par.add_run(f"{i}. {s}")


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    p2 = doc.add_paragraph("Скриншоты работы приложения")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    red_placeholder(
        doc.add_paragraph(),
        "[ЗАПОЛНИТЕ: вставьте скриншоты основных экранов клиента (все классы, выбор класса, расписание класса, полное расписание) "
        "и при необходимости сервера/логов. Каждый рисунок подпишите по ГОСТ: «Рисунок A.1 – ...».]\n",
    )

    doc.add_page_break()
    p3 = doc.add_paragraph("ПРИЛОЖЕНИЕ Б")
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.runs[0].bold = True

    p4 = doc.add_paragraph("Листинг ключевых модулей")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].bold = True
    red_placeholder(
        doc.add_paragraph(),
        "[ЗАПОЛНИТЕ: вставьте листинги (без цветного выделения) или выдержки кода: `server.py`, `timetable.py`, "
        "`data_to_jsonFile.py`, `parsing_excel.py`, `parsing_word.py`, `download_fromServer.py`.]\n",
    )


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    paths = Paths(
        root=root,
        readme=root / "README.md",
        requirements=root / "requirements.txt",
        out_docx=root / "Отчет_ВКР_Расписание.docx",
    )

    readme_text = ""
    if paths.readme.exists():
        readme_text = paths.readme.read_text(encoding="utf-8", errors="ignore")

    doc = Document()
    set_normal_style(doc)
    set_page_setup(doc)
    add_page_number_footer(doc)

    add_title_page(doc)
    add_executors_page(doc)
    add_abstract(doc)
    add_toc_placeholder(doc)
    add_terms_and_definitions(doc)
    add_abbreviations(doc)
    add_intro(doc)
    add_chapter_1_theory(doc)
    add_chapter_2_practice(doc, readme_text=readme_text)
    add_user_guide(doc)
    add_conclusion(doc)
    add_sources(doc)
    add_appendices(doc)

    try:
        doc.save(str(paths.out_docx))
        print(f"Saved: {paths.out_docx}")
    except PermissionError:
        alt = paths.out_docx.with_name(paths.out_docx.stem + "_NEW" + paths.out_docx.suffix)
        doc.save(str(alt))
        print(f"Saved (locked target, wrote new): {alt}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

